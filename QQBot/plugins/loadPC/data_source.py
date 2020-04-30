import os
import math
import prettytable as pt

from docx import Document

from constant import path_storage, pc_card, talent, max_items, items, const_row, ps_card


def find_pc(sender_id=int):
    path = os.path.join(path_storage, str(sender_id))  # storage/qq/
    file_list = os.listdir(path)
    des_file = str(sender_id) + '_C.docx'
    if des_file not in file_list:
        return False
    else:
        return True


def find_ps(sender_id=int):
    path = os.path.join(path_storage, str(sender_id))  # storage/qq/
    file_list = os.listdir(path)
    des_file = str(sender_id) + '_S.docx'
    if des_file not in file_list:
        return False
    else:
        return True


def read_pc(sender_id=int):
    path = os.path.join(path_storage, str(sender_id))  # storage/qq/
    des_file = str(sender_id) + '_C.docx'
    path = os.path.join(path, des_file)  # storage/qq/qq_C.docx
    document = Document(path)
    tables = document.tables  # 获取文件中的表格集
    table = tables[0]
    pc_card['姓名'] = table.cell(1, 1).text
    pc_card['性别'] = table.cell(1, 8).text
    pc_card['年龄'] = table.cell(1, 12).text
    pc_card['种族'] = table.cell(2, 1).text
    pc_card['职业'] = table.cell(2, 8).text
    origin_line = 4  # 天赋起始栏
    index = 1
    between_index = 0
    between = [4, 3, 4, 1]
    for i in talent:
        talent[i] = table.cell(origin_line + int((index - 1) / 12), ((index - 1) % 12) + 1).text
        index += between[between_index]
        between_index = (between_index + 1) % 4
    pc_card['天赋'] = talent
    master_line = origin_line - 1 + math.ceil(len(talent) / 4) + 1  # 推算精通所在行
    master_content = table.cell(master_line, 4).text  # 精通内容（纯文本）
    master_list = str(master_content).split('\n')
    pc_card['精通'] = master_list
    degree_line = master_line + 1  # 推算职业等级所在行
    degree = table.cell(degree_line, 6).text
    degree_develop = table.cell(degree_line, 10).text
    pc_card['职业等级'] = degree
    pc_card['熟练加值'] = degree_develop
    tricks_line = degree_line + 1  # 推算戏法所在行
    tricks_content = table.cell(tricks_line, 4).text  # 戏法内容
    tricks_list = str(tricks_content).split('\n')
    pc_card['戏法'] = tricks_list
    magic_line = tricks_line + 1  # 推算法术所在行
    magic_content = table.cell(magic_line, 4).text  # 法术内容
    magic_list = str(magic_content).split('\n')
    pc_card['法术'] = magic_list
    max_trick_and_magic_line = magic_line + 1  # 推算戏法和法术上限所在行
    max_trick = table.cell(max_trick_and_magic_line, 6).text
    max_magic = table.cell(max_trick_and_magic_line, 10).text
    pc_card['戏法上限'] = max_trick
    pc_card['法术上限'] = max_magic
    # info_row = 2 + math.ceil(len(talent) / 4) + 2  # 推算人物介绍所在行
    info_row = max_trick_and_magic_line + 2  # 推算人物介绍所在行
    pc_card['人物介绍'] = table.cell(info_row, 0).text
    print(pc_card['人物介绍'])
    cur_items = [None] * max_items
    item_row = info_row + 2  # 推算物品所在行
    index = 0  # 有效物品计数
    cur_index = 0
    i_between = [2, 6, 3, 2]
    i_between_index = 0
    num = 0
    for j in range(max_items):
        cur_items[index] = table.cell(item_row + int(cur_index / 13),
                                      cur_index % 13).text  # 0->0,0  1->0,2  2->0,5  3->0,8  4->1,0
        if cur_items[index] is not None:
            num += 1
        index += 1
        cur_index += i_between[i_between_index]
        i_between_index = (i_between_index + 1) % 4

    final_items = [None] * num
    f_index = 0
    for f_i in range(max_items):
        if cur_items[f_i] is not None:
            final_items[f_index] = cur_items[f_i]
            f_index += 1

    pc_card['物品'] = final_items
    return pc_card


def read_ps(sender_id=int):
    path = os.path.join(path_storage, str(sender_id))  # storage/qq/
    des_file = str(sender_id) + '_S.docx'
    path = os.path.join(path, des_file)  # storage/qq/qq_S.docx
    document = Document(path)
    tables = document.tables  # 获取文件中的表格集
    table = tables[0]

    ps_card['血量'] = table.cell(0, 3).text
    ps_card['血量上限'] = table.cell(0, 10).text
    ps_card['经验值'] = table.cell(1, 3).text
    ps_card['经验值上限'] = table.cell(1, 10).text
    magic_consume = [None] * 9
    magic_consume[0] = table.cell(3, 0).text
    magic_consume[1] = table.cell(3, 1).text
    magic_consume[2] = table.cell(3, 2).text
    magic_consume[3] = table.cell(3, 4).text
    magic_consume[4] = table.cell(3, 6).text
    magic_consume[5] = table.cell(3, 7).text
    magic_consume[6] = table.cell(3, 8).text
    magic_consume[7] = table.cell(3, 10).text
    magic_consume[8] = table.cell(3, 11).text

    ps_card['每环法术位'] = magic_consume
    buff_content = table.cell(4, 3).text
    buff_list = str(buff_content).split('\n')
    ps_card['Buff'] = buff_list
    debuff_content = table.cell(5, 3).text
    debuff_list = str(debuff_content).split('\n')
    ps_card['Debuff'] = debuff_list
    return ps_card


def f_format(player_card):
    global p_talent_keys
    p_talent = player_card['天赋']

    if isinstance(p_talent, dict):
        p_talent_keys = p_talent.keys()
    if not isinstance(player_card, dict):
        return
    result = '角色卡\n' + const_row + '\n'
    result += '姓名：' + player_card['姓名'] + ' \t|\t性别：' + player_card['性别'] + '\t|\t年龄：' + player_card['年龄'] + '\n' + \
              '种族：' + player_card.get('种族') + ' \t|\t职业：' + player_card.get('职业') + '\n'
    result += const_row + '\n'
    result += '天赋\n'
    result += const_row + '\n'
    now = 1
    for each in p_talent_keys:
        if now == 1:
            result += each + '：' + p_talent.get(each) + '\t|\t'
        if now == 0:
            result += each + '：' + p_talent.get(each) + '\n'
        now = (now + 1) % 2
    if now == 0:
        result += '\n'
    now = 0
    result += const_row + '\n'
    result += '精通\n'
    result += const_row + '\n'
    master_list = player_card['精通']
    for each in master_list:
        result += each + '；'
    if len(master_list) > 0:
        result = result[:-1]
        result += '\n'
    result += const_row + '\n'
    result += '职业等级：' + player_card['职业等级'] + '\t|\t熟练加值：' + player_card['熟练加值'] + '\n'
    result += const_row + '\n'
    result += '戏法\n'
    result += const_row + '\n'
    tricks_list = player_card['戏法']
    for each in tricks_list:
        result += each + '；'
    if len(tricks_list) > 0:
        result = result[:-1]
        result += '\n'
    result += const_row + '\n'
    result += '法术\n'
    result += const_row + '\n'
    magic_list = player_card['法术']
    for each in magic_list:
        result += each + '；'
    if len(master_list) > 0:
        result = result[:-1]
        result += '\n'
    result += const_row + '\n'
    result += '戏法上限：' + player_card['戏法上限'] + '\t|\t法术上限：' + player_card['法术上限'] + '\n'
    result += const_row + '\n'
    result += '人物介绍\n'
    result += const_row + '\n'
    result += player_card.get('人物介绍') + '\n'
    result += const_row + '\n'
    result += '物品\n'
    result += const_row + '\n'

    p_items = player_card['物品']
    for every in p_items:
        result += every + '；'
    if len(p_items) > 0:
        result = result[:-1]
        result += '\n'
    result += const_row + '\n'
    return result


def fs_format(player_state):
    result = '状态卡\n' + const_row + '\n'
    result += '血量：' + player_state['血量'] + ' \t|\t血量上限：' + player_state['血量上限'] + '\n' + \
              '经验值：' + player_state['经验值'] + ' \t|\t经验值上限：' + player_state['经验值上限'] + '\n'
    result += const_row + '\n'
    result += '每环法术位（1-9环）' + '\n'
    # result += const_row + '\n'
    magic_consume_list = player_state['每环法术位']
    tb = pt.PrettyTable()
    tb.header = True
    tb.field_names = ['1环', '2环', '3环', '4环', '5环', '6环', '7环', '8环', '9环']
    tb.add_row([magic_consume_list[0], magic_consume_list[1], magic_consume_list[2], magic_consume_list[3],
                magic_consume_list[4], magic_consume_list[5], magic_consume_list[6], magic_consume_list[7],
                magic_consume_list[8]])
    result += str(tb) + '\n'
    result += const_row + '\n'
    result += 'Buff\n'
    print(result)
    result += const_row + '\n'
    buff_list = player_state['Buff']
    for each in buff_list:
        result += each + '\n'
    if len(buff_list) == 0:
        result += '无\n'
    result += const_row + '\n'
    result += 'Debuff\n'
    result += const_row + '\n'
    debuff_list = player_state['Debuff']
    for each in debuff_list:
        result += each + '\n'
    if len(debuff_list) == 0:
        result += '无\n'
    result += const_row + '\n'
    return result
